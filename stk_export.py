"""
Eksport STK:
- raport DOCX — wyniki pre/post/delayed + pelna sekcja psychometryczna
  (alfa Cronbacha, SEM, analiza pozycji, trafnosc, test t, d Cohena, moc, RCI),
- surowe dane XLSX — do analiz zewnetrznych (jamovi / JASP / R).

Branding Enterprise Advisors: navy #1B4F8A, Arial.
"""
from __future__ import annotations

import io
import re
import unicodedata
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

from stk_data import STKResult, STKTest, CompetencyProfile, NeedsAssessment, CATEGORY_LABELS, LEVEL_LABELS, LEVEL_KEYS, task_priority, tasks_to_str
from stk_stats import (
    alpha_interpretation, cohen_dz_interpretation, item_analysis,
    paired_prepost_analysis, reliability_block, required_n_paired,
    subscale_correlations, validate_test,
)

EA_NAVY = RGBColor(0x1B, 0x4F, 0x8A)
EA_GOLD = RGBColor(0xF5, 0xC5, 0x18)
EA_GRAY = RGBColor(0x55, 0x55, 0x55)
FONT = "Montserrat"


# ---------------------------------------------------------------------------
# Pomocnicze formatowanie
# ---------------------------------------------------------------------------
def _style_base(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.bold = True
    run.font.color.rgb = EA_NAVY
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    p.space_before = Pt(12)


def _para(doc: Document, text: str, italic: bool = False, gray: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)
    run.italic = italic
    if gray:
        run.font.color.rgb = EA_GRAY


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = EA_NAVY
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(9)
    doc.add_paragraph()


def _set_cell_bg(cell, color_hex: str) -> None:
    """Ustawia kolor tła komórki tabeli DOCX."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _competency_level_table(doc: Document, comp: "Competency",
                             desired_level: int = 0) -> None:
    """Tabela 5 poziomów kompetencji z dopasowanymi kolumnami i podświetlonym poziomem oczekiwanym."""
    COL_W = [Cm(1.4), Cm(3.6), Cm(11.0)]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0]
    for j, h in enumerate(["Poziom", "Etykieta", "Opis behawioralny"]):
        cell = hdr.cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = EA_NAVY
        cell.width = COL_W[j]

    for lk in LEVEL_KEYS:
        row_cells = table.add_row().cells
        vals = [str(lk), LEVEL_LABELS[lk], comp.get_level_description(lk)]
        for j, val in enumerate(vals):
            row_cells[j].text = ""
            run = row_cells[j].paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(9)
            row_cells[j].width = COL_W[j]
            if lk == desired_level:
                _set_cell_bg(row_cells[j], "1B4F8A")
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    doc.add_paragraph()


def _key_tasks_table(doc: Document, tasks_list: list) -> None:
    """Tabela analizy pracy: Zadanie | Częstotliwość | Ważność | Trudność | Priorytet."""
    if not tasks_list:
        return

    COL_W = [Cm(8.0), Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.4)]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0]
    for j, h in enumerate(["Zadanie", "Częstotliwość", "Ważność", "Trudność", "Priorytet"]):
        cell = hdr.cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "1B4F8A")
        cell.width = COL_W[j]

    for task in tasks_list:
        name = task.get("name", "").strip()
        if not name:
            continue
        f = task.get("frequency", 2)
        i = task.get("importance", 2)
        d = task.get("difficulty", 2)
        prio = task_priority(i)
        row_cells = table.add_row().cells
        for j, val in enumerate([name, str(f), str(i), str(d), prio]):
            row_cells[j].text = ""
            run = row_cells[j].paragraphs[0].add_run(val)
            run.font.name = FONT
            run.font.size = Pt(9)
            run.font.bold = (j == 0)
            if j == 4:
                color = EA_NAVY if prio == "Wysoki" else EA_GRAY
                run.font.color.rgb = color
            row_cells[j].width = COL_W[j]

    doc.add_paragraph()


def _fmt(x, digits: int = 3) -> str:
    if x is None:
        return "—"
    return f"{x:.{digits}f}"


def _fmt_p(p: float | None) -> str:
    if p is None:
        return "—"
    if p < 0.001:
        return "p < .001"
    return f"p = {p:.3f}".replace("0.", ".")


# ---------------------------------------------------------------------------
# Raport glowny
# ---------------------------------------------------------------------------
def build_report_docx(test: STKTest, results: list[STKResult]) -> bytes:
    """Buduje raport DOCX: wyniki + psychometria. Zwraca bajty pliku."""
    pre = [r for r in results if r.measurement_type == "pre"]
    post = [r for r in results if r.measurement_type == "post"]
    delayed = [r for r in results if r.measurement_type == "delayed"]

    doc = Document()
    _style_base(doc)

    # --- Naglowek ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Raport pomiaru kompetencji — STK")
    run.font.name = FONT
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = EA_NAVY
    _para(doc, test.name, italic=True)
    _para(doc, f"Data raportu: {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
               f"Pytań: {len(test.questions)} | Max: {test.max_total_score} pkt | "
               f"Wyników pre: {len(pre)} | post: {len(post)} | delayed: {len(delayed)}", gray=True)

    # --- 1. Poprawnosc testu ---
    _heading(doc, "1. Poprawność struktury testu", 2)
    issues = validate_test(test)
    if issues:
        _para(doc, f"Wykryto {len(issues)} problemów z kluczem odpowiedzi:")
        for issue in issues:
            _para(doc, f"  - {issue}")
    else:
        _para(doc, "Klucz odpowiedzi poprawny: każde pytanie ma 4 opcje, dokładnie "
                   "jedna najlepsza (4 pkt) i jedna najgorsza (1 pkt).")

    # --- 2. Wyniki indywidualne ---
    _heading(doc, "2. Wyniki indywidualne", 2)
    rows = []
    for r in results:
        pct = round(r.total_score / test.max_total_score * 100, 1) if test.max_total_score else 0
        rows.append([r.participant_name, r.measurement_type.upper(),
                     f"{r.total_score}/{test.max_total_score}", f"{pct}%",
                     r.completed_at[:16].replace("T", " ")])
    _table(doc, ["Uczestnik", "Pomiar", "Wynik", "%", "Data"], rows)

    # --- 3. Rzetelnosc (alfa Cronbacha) ---
    _heading(doc, "3. Rzetelność — alfa Cronbacha", 2)
    for label, subset in (("PRE", pre), ("POST", post), ("DELAYED", delayed)):
        if not subset:
            continue
        block = reliability_block(subset, test, label)
        _heading(doc, f"Pomiar {label} (N = {block.n}, pozycji = {block.k_items})", 3)
        _table(doc, ["Wskaźnik", "Wartość", "Interpretacja"], [
            ["Alfa Cronbacha (cały test)", _fmt(block.alpha), block.alpha_note],
            ["SD wyniku łącznego", _fmt(block.sd, 2), ""],
            ["SEM (błąd standardowy pomiaru)", _fmt(block.sem, 2),
             "przedział ufności wyniku: +/- 1.96 x SEM"],
        ])
        sub_rows = [[name, str(k), _fmt(a), alpha_interpretation(a) if a is not None else "—"]
                    for name, k, a in block.subscale_alphas]
        if sub_rows:
            _para(doc, "Alfa per kompetencja (podskale):")
            _table(doc, ["Kompetencja", "Pozycji", "Alfa", "Interpretacja"], sub_rows)
        if block.n < 30:
            _para(doc, f"Uwaga: N = {block.n} — oszacowanie orientacyjne. Zalecany "
                       "pilotaż N = 20-30, badanie główne N = 100-120 (plan ITET-2).",
                  italic=True, gray=True)
    _para(doc, "Nota: dla testów sytuacyjnych (SJT) alfa bywa zaniżona, bo pozycje "
               "obejmują heterogeniczne konteksty (Catano et al. 2012). Przy niskiej "
               "alfie zalecana ocena rzetelności metodą test-retest.", italic=True, gray=True)

    # --- 4. Analiza pozycji ---
    if pre:
        analysis_set, set_label = pre, "PRE"
    elif post:
        analysis_set, set_label = post, "POST"
    else:
        analysis_set, set_label = delayed, "DELAYED"
    if analysis_set:
        _heading(doc, "4. Analiza pozycji (trudność i moc dyskryminacyjna)", 2)
        _para(doc, f"Na podstawie pomiaru {set_label} "
                   f"(N = {len(analysis_set)}). Trudność = średni wynik pozycji / 2 "
                   "(wyższa wartość = łatwiejsza pozycja). Dyskryminacja = skorygowana "
                   "korelacja pozycja-wynik; wartości < 0.20 wskazują słabą pozycję (Ebel).")
        items = item_analysis(analysis_set, test)
        _table(doc, ["Pytanie", "Kompetencja", "Trudność", "Dyskryminacja", "Uwagi"], [
            [it.question_id, it.competency_name, _fmt(it.difficulty),
             _fmt(it.discrimination), it.flag or "OK"]
            for it in items
        ])

    # --- 5. Trafnosc ---
    _heading(doc, "5. Trafność", 2)
    _para(doc, "Trafność treściowa: scenariusze i klucz ekspercki generowane wg formatu "
               "najlepsze + najgorsze zachowanie (Filipowicz 2014, s. 41; Prokopowicz "
               "et al. 2014), na podstawie behawioralnych opisów poziomów 1-5. Klucz "
               "wymaga przeglądu eksperckiego przez trenera przed użyciem.")
    if analysis_set and len(analysis_set) >= 3 and len(test.competency_names) >= 2:
        _para(doc, "Trafność teoretyczna (struktura wewnętrzna): korelacje Spearmana "
                   "między podskalami. Umiarkowane korelacje dodatnie (~0.2-0.6) "
                   "wspierają odrębność, a zarazem pokrewność konstruktów.")
        corr_rows = [[a, b, _fmt(rho)] for a, b, rho in
                     subscale_correlations(analysis_set, test)]
        _table(doc, ["Podskala A", "Podskala B", "rho Spearmana"], corr_rows)

    # --- 6. Zmiana miedzy pomiarami: test t, d Cohena, moc, RCI ---
    pair_defs = [
        ("PRE → POST (efekt szkolenia)", "PRE", "POST", pre, post),
        ("POST → DELAYED (utrzymanie efektu)", "POST", "DELAYED", post, delayed),
        ("PRE → DELAYED (efekt trwały)", "PRE", "DELAYED", pre, delayed),
    ]
    available_pairs = [p for p in pair_defs if p[3] and p[4]]
    if available_pairs:
        _heading(doc, "6. Zmiana między pomiarami — istotność, wielkość efektu, moc", 2)
        for pair_label, lbl_a, lbl_b, set_a, set_b in available_pairs:
            pa = paired_prepost_analysis(set_a, set_b, test)
            _heading(doc, pair_label, 3)
            if pa.n == 0:
                _para(doc, "Brak sparowanych wyników (ten sam uczestnik w obu pomiarach).")
                continue
            _table(doc, ["Wskaźnik", "Wartość", "Interpretacja"], [
                ["N par", str(pa.n), ""],
                [f"Średnia {lbl_a} (SD)", f"{_fmt(pa.mean_pre, 2)} ({_fmt(pa.sd_pre, 2)})", ""],
                [f"Średnia {lbl_b} (SD)", f"{_fmt(pa.mean_post, 2)} ({_fmt(pa.sd_post, 2)})", ""],
                ["Średnia zmiana", _fmt(pa.mean_diff, 2), ""],
                ["Test t dla prób zależnych",
                 f"t({pa.df}) = {_fmt(pa.t_stat, 2)}" if pa.t_stat is not None else "—",
                 _fmt_p(pa.p_value)],
                ["d Cohena (dz)", _fmt(pa.cohen_dz, 2),
                 cohen_dz_interpretation(pa.cohen_dz) if pa.cohen_dz is not None else "—"],
                ["Moc post-hoc (alfa = .05)", _fmt(pa.power_posthoc, 2),
                 "aproksymacja normalna (Cohen 1988)"],
            ])
            if pa.rci_per_participant:
                _para(doc, "RCI — wskaźnik rzetelnej zmiany (Jacobson i Truax 1991): "
                           f"|RCI| > 1.96 oznacza zmianę większą niż błąd pomiaru "
                           f"(SEM {lbl_a} = {_fmt(pa.sem_pre, 2)}, "
                           f"alfa {lbl_a} = {_fmt(pa.alpha_pre)}).")
                _table(doc, ["Uczestnik", "RCI", "Ocena"],
                       [[n, str(v), o] for n, v, o in pa.rci_per_participant])
            elif pa.alpha_pre is None:
                _para(doc, f"RCI niedostępny: alfa Cronbacha pomiaru {lbl_a} nieobliczalna "
                           "(za mało uczestników lub brak zróżnicowania wyników).",
                      italic=True, gray=True)
        _para(doc, "Wymagane N par dla mocy 0.80 (dwustronnie, alfa = .05): "
                   f"d = 0.2 -> N = {required_n_paired(0.2)}; "
                   f"d = 0.5 -> N = {required_n_paired(0.5)}; "
                   f"d = 0.71 -> N = {required_n_paired(0.71)} (Rogge et al. 2022); "
                   f"d = 0.8 -> N = {required_n_paired(0.8)}.")

    # --- 7. Nota metodyczna ---
    _heading(doc, "7. Nota metodyczna", 2)
    _para(doc, "Instrument: Sytuacyjny Test Kompetencyjny (SJT), format najlepsze + "
               "najgorsze zachowanie, scoring 0-1-2 pkt na pytanie. Pomiar zmiany "
               "kompetencji w schemacie pre → post → delayed (pomiar odroczony) "
               "zgodnie z LTEM Tier 5 (Thalheimer) i planem badawczym modelu ITET-2 "
               "(PB2). Korelacje dla zmiennych porządkowych "
               "liczone współczynnikiem rho Spearmana (Czakon 2020). Istotność zmiany: "
               "test t dla prób zależnych; wielkość efektu: d Cohena dla prób "
               "zależnych (dz); zmiana indywidualna: RCI (Jacobson i Truax 1991).",
          gray=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# Surowe dane XLSX — do analiz w jamovi / JASP / R
# ===========================================================================
MEAS_ORDER = ["pre", "post", "delayed"]

_PL_MAP = str.maketrans({"ł": "l", "Ł": "L"})


def _safe_name(name: str) -> str:
    """Nazwa kolumny przyjazna dla jamovi/R: ASCII, snake_case."""
    s = name.translate(_PL_MAP)
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_").lower()
    return s or "kol"


def _add_sheet(wb: Workbook, title: str, headers: list[str],
               rows: list[list], col_width: int = 14) -> None:
    ws = wb.create_sheet(title)
    ws.append(headers)
    fill = PatternFill(start_color="1B4F8A", end_color="1B4F8A", fill_type="solid")
    font = Font(color="FFFFFF", bold=True, name="Arial", size=10)
    for j in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=j)
        cell.fill = fill
        cell.font = font
        ws.column_dimensions[get_column_letter(j)].width = col_width
    for row in rows:
        ws.append(row)
    ws.freeze_panes = "A2"


def build_raw_xlsx(test: STKTest, results: list[STKResult]) -> bytes:
    """Surowe wyniki STK w formatach gotowych pod jamovi/JASP/R.

    Arkusze:
    - dane_szerokie: 1 wiersz = uczestnik (kolumny pre/post/delayed) — test t dla prob zaleznych
    - pozycje: 1 wiersz = uczestnik x pomiar, kolumny Q1..Qn (0-2) — alfa Cronbacha
    - dane_dlugie: uczestnik x pomiar x kompetencja — eksploracja/wykresy
    - odpowiedzi: wybory best/worst per pytanie — analiza dystraktorow
    - klucz_testu: scenariusze, opcje, klucz
    - slownik: opis zmiennych + instrukcja analiz w jamovi
    """
    wb = Workbook()
    wb.remove(wb.active)

    comp_names = test.competency_names
    comp_cols = {name: _safe_name(name) for name in comp_names}
    q_ids = [q.id for q in test.questions]

    # ostatni wynik per uczestnik x pomiar (jak w dashboardzie)
    by_participant: dict[str, dict[str, STKResult]] = {}
    for r in results:
        by_participant.setdefault(r.participant_id, {})[r.measurement_type] = r

    # --- 1. dane_szerokie ---
    headers = ["id", "uczestnik"]
    for m in MEAS_ORDER:
        headers += [f"wynik_{m}", f"pct_{m}"]
    for name in comp_names:
        for m in MEAS_ORDER:
            headers.append(f"{comp_cols[name]}_{m}")
    rows = []
    for pid, meas in by_participant.items():
        any_r = next(iter(meas.values()))
        row: list = [pid, any_r.participant_name]
        for m in MEAS_ORDER:
            if m in meas:
                total = meas[m].total_score
                pct = round(total / test.max_total_score * 100, 1) if test.max_total_score else None
                row += [total, pct]
            else:
                row += [None, None]
        by_comp = {m: meas[m].score_by_competency(test) for m in meas}
        for name in comp_names:
            for m in MEAS_ORDER:
                row.append(by_comp[m][name]["score"] if m in meas else None)
        rows.append(row)
    _add_sheet(wb, "dane_szerokie", headers, rows)

    # --- 2. pozycje (macierz do alfy Cronbacha) ---
    headers = ["id", "uczestnik", "pomiar", "data"] + q_ids + ["wynik_laczny"]
    rows = []
    for r in results:
        by_id = {a.question_id: a.score for a in r.answers}
        rows.append([r.participant_id, r.participant_name, r.measurement_type,
                     r.completed_at[:16].replace("T", " ")]
                    + [by_id.get(qid) for qid in q_ids]
                    + [r.total_score])
    _add_sheet(wb, "pozycje", headers, rows, col_width=10)

    # --- 3. dane_dlugie ---
    headers = ["id", "uczestnik", "pomiar", "kompetencja", "wynik", "max", "pct"]
    rows = []
    for r in results:
        by_comp = r.score_by_competency(test)
        for name in comp_names:
            v = by_comp[name]
            rows.append([r.participant_id, r.participant_name, r.measurement_type,
                         name, v["score"], v["max"], v["pct"]])
    _add_sheet(wb, "dane_dlugie", headers, rows)

    # --- 4. odpowiedzi (dystraktory) ---
    q_by_id = {q.id: q for q in test.questions}
    headers = ["id", "uczestnik", "pomiar", "pytanie", "kompetencja",
               "wybor_najlepsza", "wybor_najgorsza",
               "klucz_najlepsza", "klucz_najgorsza", "punkty"]
    rows = []
    for r in results:
        for a in r.answers:
            q = q_by_id.get(a.question_id)
            if q is None or not q.options:
                continue
            kb = max(range(len(q.options)), key=lambda i: q.options[i].score)
            kw = min(range(len(q.options)), key=lambda i: q.options[i].score)
            rows.append([
                r.participant_id, r.participant_name, r.measurement_type,
                a.question_id, q.competency_name,
                chr(65 + a.best_index) if a.best_index >= 0 else None,
                chr(65 + a.worst_index) if a.worst_index >= 0 else None,
                chr(65 + kb), chr(65 + kw), a.score,
            ])
    _add_sheet(wb, "odpowiedzi", headers, rows)

    # --- 5. klucz_testu ---
    headers = ["pytanie", "kompetencja", "scenariusz",
               "opcja_A", "opcja_B", "opcja_C", "opcja_D",
               "pkt_A", "pkt_B", "pkt_C", "pkt_D"]
    rows = []
    for q in test.questions:
        texts = [o.text for o in q.options] + [None] * (4 - len(q.options))
        scores = [o.score for o in q.options] + [None] * (4 - len(q.options))
        rows.append([q.id, q.competency_name, q.scenario] + texts[:4] + scores[:4])
    _add_sheet(wb, "klucz_testu", headers, rows, col_width=30)

    # --- 6. slownik ---
    headers = ["arkusz / zmienna", "opis"]
    rows = [
        ["dane_szerokie", "1 wiersz = uczestnik; kolumny wynik_pre/post/delayed oraz "
                          "wyniki podskal per pomiar. Puste komórki = brak pomiaru (NA)."],
        ["pozycje", "1 wiersz = uczestnik x pomiar; kolumny Q1..Qn = punkty pozycji (0-2)."],
        ["dane_dlugie", "Format długi: uczestnik x pomiar x kompetencja."],
        ["odpowiedzi", "Wybory uczestnika (A-D) vs klucz — analiza dystraktorów."],
        ["klucz_testu", "Scenariusze, treści opcji i punktacja (4=najlepsza, 1=najgorsza)."],
        ["punkty pozycji", "0 = brak trafień, 1 = jedno trafienie, 2 = oba (best+worst), "
                           "wg Filipowicza (2014, s. 41)."],
        ["", ""],
        ["JAMOVI — alfa Cronbacha", "Arkusz 'pozycje': przefiltruj wiersze po kolumnie "
                                    "'pomiar' (np. pre), potem Analyses > Factor > "
                                    "Reliability Analysis, zmienne Q1..Qn."],
        ["JAMOVI — test t (pre/post)", "Arkusz 'dane_szerokie': Analyses > T-Tests > "
                                       "Paired Samples T-Test, pary wynik_pre + wynik_post "
                                       "(analogicznie wynik_delayed). Wielkość efektu: "
                                       "zaznacz Effect size (d Cohena)."],
        ["JAMOVI — Spearman", "Arkusz 'dane_szerokie': Analyses > Regression > Correlation "
                              "Matrix, zaznacz Spearman (zmienne porządkowe, konwencja "
                              "ITET-2/Czakon 2020)."],
        ["Eksport", f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
                    f"Test: {test.name} | Pytań: {len(test.questions)} | "
                    f"Wyników: {len(results)}"],
    ]
    _add_sheet(wb, "slownik", headers, rows, col_width=40)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Zbiorczy raport — wiele profili (mapa kompetencji firmy)
# ---------------------------------------------------------------------------
def _radar_chart_png(
    profile: CompetencyProfile,
    needs: "NeedsAssessment | None",
) -> bytes | None:
    """Wykres radarowy PNG — poziomy kompetencji. Zwraca None przy błędzie lub za mało danych."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np
    except ImportError:
        return None

    comp_names = [c.name for c in profile.competencies][:12]
    N = len(comp_names)
    if N < 3:
        return None

    needs_map: dict[str, tuple[int, int]] = {}
    if needs:
        for item in needs.items:
            needs_map[item.competency_name] = (item.current_level, item.desired_level)

    desired = [needs_map.get(n, (3, 3))[1] for n in comp_names]
    has_current = needs and any(needs_map.get(n, (0, 0))[0] > 0 for n in comp_names)
    current = [needs_map.get(n, (0, 0))[0] for n in comp_names] if has_current else None

    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
    angles_c = angles + angles[:1]
    desired_c = desired + desired[:1]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw={"polar": True})
    _NAVY = "#1B4F8A"
    _GOLD = "#F5C518"

    def _lbl(name: str) -> str:
        """Krótka etykieta: pierwsze słowo przed '/' lub '/', max 12 znaków."""
        main = name.split("/")[0].strip()
        words = main.split()
        if not words:
            return name[:12]
        first = words[0]
        if len(words) == 1 or len(first) > 12:
            return first[:13]
        combined = f"{first}\n{words[1]}"
        return combined if len(words[1]) <= 12 else first

    ax.set_ylim(0, 5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(["1", "2", "3", "4", "5"], color="#aaaaaa", size=7)
    ax.set_xticks(angles)
    ax.set_xticklabels([_lbl(n) for n in comp_names], size=7.5, color="#222222",
                       multialignment="center")
    ax.tick_params(pad=18)
    ax.grid(color="#dddddd", linewidth=0.5)
    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")

    # Poziom oczekiwany — granat z markerami
    ax.fill(angles_c, desired_c, color=_NAVY, alpha=0.2)
    ax.plot(angles_c, desired_c, color=_NAVY, linewidth=2.5,
            marker="o", markersize=6, markerfacecolor=_NAVY,
            label="Poziom oczekiwany")

    # Poziom aktualny — złoto z markerami (jeśli dostępny)
    if current:
        current_c = current + current[:1]
        ax.fill(angles_c, current_c, color=_GOLD, alpha=0.15)
        ax.plot(angles_c, current_c, color=_GOLD, linewidth=2, linestyle="--",
                marker="o", markersize=5, markerfacecolor=_GOLD,
                label="Poziom aktualny")

    # Linia normy poziom 3
    norm_c = [3] * (N + 1)
    ax.plot(angles_c, norm_c, color="#cccccc", linewidth=0.8, linestyle=":", label="Norma (3)")

    # Wartości liczbowe przy markerach (poziom oczekiwany)
    for angle, val in zip(angles, desired):
        ax.text(angle, val + 0.3, str(val), ha="center", va="bottom",
                size=7, color=_NAVY, fontweight="bold")

    ax.legend(loc="upper left", bbox_to_anchor=(-0.18, -0.06), ncol=3,
               fontsize=7.5, frameon=False)
    ax.set_title(profile.company.position_name, size=10, color=_NAVY, fontweight="bold", pad=30)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def build_combined_docx(profiles: list[CompetencyProfile],
                        needs: "list[NeedsAssessment | None] | None" = None) -> bytes:
    """Generuje zbiorczy DOCX z profilami kompetencji wszystkich stanowisk.

    needs: lista NeedsAssessment (lub None) równoległa do profiles — do wykresu radarowego.
    """
    _needs = needs or [None] * len(profiles)

    doc = Document()
    _style_base(doc)

    # --- Strona tytułowa ---
    doc.add_paragraph()
    title = doc.add_paragraph()
    run = title.add_run("MAPA KOMPETENCJI")
    run.font.name = FONT
    run.font.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = EA_NAVY
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Złota linia dekoracyjna (tabela 1×1 z tłem złotym)
    gold_tbl = doc.add_table(rows=1, cols=1)
    cell = gold_tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5C518")
    tcPr.append(shd)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    run2 = subtitle.add_run(f"Zbiorczy profil kompetencji — {len(profiles)} stanowisk")
    run2.font.name = FONT
    run2.font.size = Pt(14)
    run2.font.color.rgb = EA_GRAY
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    gen_p = doc.add_paragraph()
    run3 = gen_p.add_run(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run3.font.name = FONT
    run3.font.size = Pt(9)
    run3.font.color.rgb = EA_GRAY
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ea_p = doc.add_paragraph()
    run4 = ea_p.add_run("Enterprise Advisors")
    run4.font.name = FONT
    run4.font.size = Pt(10)
    run4.font.bold = True
    run4.font.color.rgb = EA_NAVY
    ea_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Spis stanowisk ---
    _heading(doc, "Spis stanowisk", level=1)
    for i, p in enumerate(profiles, 1):
        row_p = doc.add_paragraph()
        num_run = row_p.add_run(f"{i}.  ")
        num_run.font.name = FONT
        num_run.font.bold = True
        num_run.font.color.rgb = EA_GOLD
        num_run.font.size = Pt(10)
        name_run = row_p.add_run(f"{p.company.position_name}")
        name_run.font.name = FONT
        name_run.font.bold = True
        name_run.font.size = Pt(10)
        detail_run = row_p.add_run(f"  —  {p.company.company_name}  ({len(p.competencies)} kompetencji)")
        detail_run.font.name = FONT
        detail_run.font.size = Pt(10)
        detail_run.font.color.rgb = EA_GRAY
    doc.add_page_break()

    # --- Profil per stanowisko ---
    for i, (profile, na) in enumerate(zip(profiles, _needs), 1):
        # Mapa poziomów oczekiwanych per kompetencja (z NeedsAssessment)
        desired_map: dict[str, int] = {}
        if na:
            for item in na.items:
                desired_map[item.competency_name] = item.desired_level
        c = profile.company

        # Nagłówek stanowiska z numerem w złocie
        h_p = doc.add_paragraph()
        num_r = h_p.add_run(f"{i}.  ")
        num_r.font.name = FONT
        num_r.font.bold = True
        num_r.font.size = Pt(18)
        num_r.font.color.rgb = EA_GOLD
        pos_r = h_p.add_run(c.position_name)
        pos_r.font.name = FONT
        pos_r.font.bold = True
        pos_r.font.size = Pt(18)
        pos_r.font.color.rgb = EA_NAVY

        _para(doc, f"{c.company_name}  |  {c.industry}  |  {c.size}  |  Poziom: {c.position_level}", gray=True)
        if c.key_tasks_list:
            _heading(doc, "Analiza pracy — kluczowe zadania", level=3)
            _para(doc, "Zadania oceniane w skali 1-3. Zadania o wysokiej ważności i trudności wyznaczają kompetencje kluczowe.", gray=True)
            _key_tasks_table(doc, c.key_tasks_list)
        elif c.key_tasks:
            _para(doc, f"Kluczowe zadania: {c.key_tasks}", italic=True)

        # Wykres radarowy
        png = _radar_chart_png(profile, na)
        if png:
            doc.add_paragraph()
            chart_p = doc.add_paragraph()
            chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_chart = chart_p.add_run()
            run_chart.add_picture(io.BytesIO(png), width=Inches(4.5))

        doc.add_paragraph()

        # Kompetencje
        for j, comp in enumerate(profile.competencies, 1):
            cat = CATEGORY_LABELS.get(comp.category, comp.category)
            comp_p = doc.add_paragraph()
            comp_num = comp_p.add_run(f"{j}. ")
            comp_num.font.name = FONT
            comp_num.font.bold = True
            comp_num.font.size = Pt(11)
            comp_num.font.color.rgb = EA_GOLD
            comp_name = comp_p.add_run(f"{comp.name}")
            comp_name.font.name = FONT
            comp_name.font.bold = True
            comp_name.font.size = Pt(11)
            comp_name.font.color.rgb = EA_NAVY
            comp_cat = comp_p.add_run(f"  [{cat}]")
            comp_cat.font.name = FONT
            comp_cat.font.size = Pt(9)
            comp_cat.font.color.rgb = EA_GRAY

            _para(doc, f"Definicja: {comp.definition}")

            if comp.indicators:
                p_ind = doc.add_paragraph()
                run_ind = p_ind.add_run("Wskaźniki behawioralne:")
                run_ind.font.name = FONT
                run_ind.font.bold = True
                run_ind.font.size = Pt(10)
                for ind in comp.indicators:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.add_run(ind).font.name = FONT

            desired_lv = desired_map.get(comp.name, 0)
            _competency_level_table(doc, comp, desired_lv)

        if i < len(profiles):
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_profile_docx(profile: CompetencyProfile,
                       assessment: "NeedsAssessment | None" = None) -> bytes:
    """Indywidualny raport DOCX dla jednego uczestnika — profil kompetencji + radar + analiza potrzeb."""
    doc = Document()
    _style_base(doc)

    # --- Strona tytułowa ---
    doc.add_paragraph()
    title = doc.add_paragraph()
    r = title.add_run("PROFIL KOMPETENCJI")
    r.font.name = FONT; r.font.bold = True; r.font.size = Pt(26); r.font.color.rgb = EA_NAVY
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    gold_tbl = doc.add_table(rows=1, cols=1)
    cell = gold_tbl.rows[0].cells[0]
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F5C518")
    tcPr.append(shd)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

    pos_p = doc.add_paragraph()
    pos_r = pos_p.add_run(profile.company.position_name)
    pos_r.font.name = FONT; pos_r.font.bold = True; pos_r.font.size = Pt(16); pos_r.font.color.rgb = EA_NAVY
    pos_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run(f"{profile.company.company_name}  |  {profile.company.industry}")
    sub_r.font.name = FONT; sub_r.font.size = Pt(12); sub_r.font.color.rgb = EA_GRAY
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    gen_p = doc.add_paragraph()
    gen_r = gen_p.add_run(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Enterprise Advisors")
    gen_r.font.name = FONT; gen_r.font.size = Pt(9); gen_r.font.color.rgb = EA_GRAY
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Wykres radarowy ---
    png = _radar_chart_png(profile, assessment)
    if png:
        doc.add_paragraph()
        chart_p = doc.add_paragraph()
        chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_p.add_run().add_picture(io.BytesIO(png), width=Inches(4.5))
    doc.add_page_break()

    # --- Analiza pracy ---
    if profile.company.key_tasks_list:
        _heading(doc, "Analiza pracy — kluczowe zadania", level=1)
        _para(doc, "Zadania oceniane w skali 1-3. Zadania o wysokiej ważności i trudności wyznaczają kompetencje kluczowe.", gray=True)
        _key_tasks_table(doc, profile.company.key_tasks_list)
    elif profile.company.key_tasks:
        _heading(doc, "Kluczowe zadania", level=1)
        _para(doc, profile.company.key_tasks, italic=True)

    # --- Kompetencje ---
    desired_map: dict[str, int] = {}
    if assessment:
        for item in assessment.items:
            desired_map[item.competency_name] = item.desired_level

    _heading(doc, "Profil kompetencji", level=1)
    for j, comp in enumerate(profile.competencies, 1):
        cat = CATEGORY_LABELS.get(comp.category, comp.category)
        comp_p = doc.add_paragraph()
        num_r2 = comp_p.add_run(f"{j}. ")
        num_r2.font.name = FONT; num_r2.font.bold = True; num_r2.font.size = Pt(11); num_r2.font.color.rgb = EA_GOLD
        name_r = comp_p.add_run(comp.name)
        name_r.font.name = FONT; name_r.font.bold = True; name_r.font.size = Pt(11); name_r.font.color.rgb = EA_NAVY
        cat_r = comp_p.add_run(f"  [{cat}]")
        cat_r.font.name = FONT; cat_r.font.size = Pt(9); cat_r.font.color.rgb = EA_GRAY

        _para(doc, f"Definicja: {comp.definition}")
        if comp.indicators:
            pi = doc.add_paragraph()
            ri = pi.add_run("Wskaźniki behawioralne:")
            ri.font.name = FONT; ri.font.bold = True; ri.font.size = Pt(10)
            for ind in comp.indicators:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(ind).font.name = FONT

        _competency_level_table(doc, comp, desired_map.get(comp.name, 0))

    # --- Analiza potrzeb ---
    if assessment and assessment.items:
        doc.add_page_break()
        _heading(doc, "Analiza potrzeb szkoleniowych", level=1)
        sorted_items = sorted(assessment.items, key=lambda x: x.priority_score, reverse=True)
        rows = []
        for item in sorted_items:
            gap_txt = str(item.gap) if item.gap else "0"
            rows.append([
                item.competency_name,
                f"{item.current_level} — {LEVEL_LABELS.get(item.current_level, '')}",
                f"{item.desired_level} — {LEVEL_LABELS.get(item.desired_level, '')}",
                gap_txt,
                str(item.importance),
                str(item.priority_score),
            ])
        _table(doc, ["Kompetencja", "Aktualny poziom", "Poziom oczekiwany", "Luka", "Waga", "Priorytet"], rows)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_incidents_docx(
    incidents: list,
    position_name: str = "",
    company_name: str = "",
) -> bytes:
    """Generuje DOCX z incydentami krytycznymi — sformatowany raport dla trenera."""
    from stk_data import CriticalIncident
    incs = [
        (i if isinstance(i, CriticalIncident) else CriticalIncident.from_dict(i))
        for i in incidents
    ]

    doc = Document()
    _style_base(doc)

    doc.add_paragraph()
    title = doc.add_paragraph()
    run = title.add_run("INCYDENTY KRYTYCZNE")
    run.font.name = FONT; run.font.bold = True
    run.font.size = Pt(24); run.font.color.rgb = EA_NAVY
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    gold_tbl = doc.add_table(rows=1, cols=1)
    c = gold_tbl.rows[0].cells[0]
    tcPr = c._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5C518")
    tcPr.append(shd)
    c.paragraphs[0].paragraph_format.space_before = Pt(2)
    c.paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

    if position_name or company_name:
        sub = doc.add_paragraph()
        sub_run = sub.add_run(position_name + (f"  —  {company_name}" if company_name else ""))
        sub_run.font.name = FONT; sub_run.font.size = Pt(13); sub_run.font.color.rgb = EA_GRAY
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    gen_p = doc.add_paragraph()
    gr = gen_p.add_run(
        f"Liczba incydentów: {len(incs)}  |  Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    gr.font.name = FONT; gr.font.size = Pt(9); gr.font.color.rgb = EA_GRAY
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ea_p = doc.add_paragraph()
    ear = ea_p.add_run("Enterprise Advisors")
    ear.font.name = FONT; ear.font.size = Pt(10)
    ear.font.bold = True; ear.font.color.rgb = EA_NAVY
    ea_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    _heading(doc, "Nota metodyczna", level=1)
    nota = doc.add_paragraph()
    nota_run = nota.add_run(
        "Poniższe incydenty krytyczne zebrano od uczestników warsztatu (Prokopowicz et al. 2014). "
        "Stanowią one surowiec do tworzenia dylematów sytuacyjnych w teście STK. "
        "Przy konwersji: odanominizuj szczegóły identyfikujące, zachowaj istotę sytuacji i charakter decyzji. "
        "Najlepsza i najgorsza reakcja z incydentu mogą służyć jako opcje scoringowe (wynik 4 i 1 pkt)."
    )
    nota_run.font.name = FONT; nota_run.font.size = Pt(10)
    doc.add_paragraph()

    by_comp: dict[str, list] = {}
    for inc in incs:
        by_comp.setdefault(inc.competency_name or "Nieokreślona", []).append(inc)

    FIELDS = [
        ("situation",         "Sytuacja (co, kiedy, gdzie, kontekst)"),
        ("actors",            "Osoby zaangażowane"),
        ("action",            "Decyzja / działanie"),
        ("reasoning",         "Powód decyzji"),
        ("result",            "Rezultat"),
        ("best_alternative",  "Najlepsza możliwa reakcja"),
        ("worst_alternative", "Najgorsza możliwa reakcja"),
    ]

    for comp_name, comp_incs in by_comp.items():
        _heading(doc, comp_name, level=1)
        for n, inc in enumerate(comp_incs, 1):
            h2p = doc.add_paragraph()
            h2r = h2p.add_run(f"Incydent {n}")
            h2r.font.name = FONT; h2r.font.bold = True
            h2r.font.size = Pt(11); h2r.font.color.rgb = EA_NAVY

            tbl = doc.add_table(rows=len(FIELDS), cols=2)
            tbl.style = "Table Grid"
            for row_i, (field_key, field_label) in enumerate(FIELDS):
                cells = tbl.rows[row_i].cells
                lbl_r = cells[0].paragraphs[0].add_run(field_label)
                lbl_r.font.name = FONT; lbl_r.font.bold = True
                lbl_r.font.size = Pt(9); lbl_r.font.color.rgb = EA_NAVY
                _set_cell_bg(cells[0], "E8F0FB")
                cells[0].width = Cm(5.5)
                val = getattr(inc, field_key, "") or ""
                val_r = cells[1].paragraphs[0].add_run(val)
                val_r.font.name = FONT; val_r.font.size = Pt(9)
                cells[1].width = Cm(11.0)
                if field_key in ("best_alternative", "worst_alternative"):
                    _set_cell_bg(cells[0], "FFF3CD")
            doc.add_paragraph()
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_incidents_xlsx(
    incidents: list,
    position_name: str = "",
    company_name: str = "",
) -> bytes:
    """Generuje XLSX z incydentami krytycznymi."""
    from stk_data import CriticalIncident
    from openpyxl.styles import Alignment as XlAlign

    incs = [
        (i if isinstance(i, CriticalIncident) else CriticalIncident.from_dict(i))
        for i in incidents
    ]

    wb = Workbook()
    ws = wb.active
    ws.title = "Incydenty_krytyczne"
    NAVY = "1B4F8A"
    LIGHT = "E8F0FB"
    GOLD_L = "FFF3CD"

    ws.merge_cells("A1:I1")
    th = ws.cell(row=1, column=1,
                 value=f"Incydenty krytyczne — {position_name}" + (f" ({company_name})" if company_name else ""))
    th.font = Font(name=FONT, bold=True, size=13, color="FFFFFF")
    th.fill = PatternFill("solid", fgColor=NAVY)
    th.alignment = XlAlign(horizontal="center")
    ws.row_dimensions[1].height = 24

    headers = [
        "Lp.", "Kompetencja", "Sytuacja", "Osoby zaangażowane",
        "Decyzja / działanie", "Powód decyzji", "Rezultat",
        "Najlepsza reakcja", "Najgorsza reakcja",
    ]
    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=2, column=ci, value=h)
        c.font = Font(name=FONT, bold=True, color="FFFFFF", size=9)
        c.fill = PatternFill("solid", fgColor=NAVY)
        c.alignment = XlAlign(wrap_text=True, vertical="top")

    for ri, inc in enumerate(incs, 1):
        bg = LIGHT if ri % 2 == 0 else None
        for ci, val in enumerate([
            ri, inc.competency_name, inc.situation, inc.actors,
            inc.action, inc.reasoning, inc.result,
            inc.best_alternative, inc.worst_alternative,
        ], 1):
            cell = ws.cell(row=ri + 2, column=ci, value=val)
            cell.font = Font(name=FONT, size=9)
            cell.alignment = XlAlign(wrap_text=True, vertical="top")
            if ci in (8, 9):
                cell.fill = PatternFill("solid", fgColor=GOLD_L)
            elif bg:
                cell.fill = PatternFill("solid", fgColor=bg)
        ws.row_dimensions[ri + 2].height = 80

    for ci, w in enumerate([4, 22, 36, 20, 30, 25, 25, 30, 30], 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_combined_xlsx(profiles: list[CompetencyProfile]) -> bytes:
    """Generuje zbiorczy XLSX — jeden arkusz per stanowisko + arkusz podsumowania."""
    wb = Workbook()
    _NAVY_HEX = "1B4F8A"
    _LIGHT_HEX = "E8F0FB"

    def _hdr(ws, row, col, text):
        c = ws.cell(row=row, column=col, value=text)
        c.font = Font(name=FONT, bold=True, color="FFFFFF", size=10)
        c.fill = PatternFill("solid", fgColor=_NAVY_HEX)
        from openpyxl.styles import Alignment
        c.alignment = Alignment(wrap_text=True, vertical="top")

    def _cell(ws, row, col, value, bold=False, bg=None):
        c = ws.cell(row=row, column=col, value=value)
        c.font = Font(name=FONT, bold=bold, size=10)
        if bg:
            c.fill = PatternFill("solid", fgColor=bg)
        from openpyxl.styles import Alignment
        c.alignment = Alignment(wrap_text=True, vertical="top")

    # ---- Arkusz Podsumowanie ----
    ws0 = wb.active
    ws0.title = "Podsumowanie"
    ws0.merge_cells("A1:F1")
    t = ws0.cell(row=1, column=1, value="Mapa Kompetencji — Podsumowanie stanowisk")
    t.font = Font(name=FONT, bold=True, size=14, color="FFFFFF")
    t.fill = PatternFill("solid", fgColor=_NAVY_HEX)
    from openpyxl.styles import Alignment
    t.alignment = Alignment(horizontal="center")

    for col, h in enumerate(["Lp.", "Stanowisko", "Firma", "Branża", "Poziom", "Liczba kompetencji"], 1):
        _hdr(ws0, 2, col, h)
    for i, p in enumerate(profiles, 1):
        bg = _LIGHT_HEX if i % 2 == 0 else None
        _cell(ws0, i + 2, 1, i, bg=bg)
        _cell(ws0, i + 2, 2, p.company.position_name, bold=True, bg=bg)
        _cell(ws0, i + 2, 3, p.company.company_name, bg=bg)
        _cell(ws0, i + 2, 4, p.company.industry, bg=bg)
        _cell(ws0, i + 2, 5, p.company.position_level, bg=bg)
        _cell(ws0, i + 2, 6, len(p.competencies), bg=bg)
    for col, w in enumerate([4, 28, 22, 18, 16, 12], 1):
        ws0.column_dimensions[get_column_letter(col)].width = w

    # ---- Arkusz per stanowisko ----
    for profile in profiles:
        # Skróć nazwę arkusza do 31 znaków (limit Excela)
        sheet_name = profile.company.position_name[:28].strip()
        # Usuń znaki niedozwolone w nazwie arkusza
        for ch in r'\/*?:[]':
            sheet_name = sheet_name.replace(ch, "-")
        ws = wb.create_sheet(sheet_name)

        ws.merge_cells("A1:I1")
        t2 = ws.cell(row=1, column=1,
                     value=f"{profile.company.position_name} — {profile.company.company_name}")
        t2.font = Font(name=FONT, bold=True, size=13, color="FFFFFF")
        t2.fill = PatternFill("solid", fgColor=_NAVY_HEX)
        t2.alignment = Alignment(horizontal="center")

        info = [
            ("Firma", profile.company.company_name),
            ("Branża", profile.company.industry),
            ("Wielkość", profile.company.size),
            ("Poziom", profile.company.position_level),
            ("Kluczowe zadania", tasks_to_str(profile.company.key_tasks_list) if profile.company.key_tasks_list else profile.company.key_tasks),
        ]
        for ri, (label, val) in enumerate(info, start=2):
            _cell(ws, ri, 1, label, bold=True, bg=_LIGHT_HEX)
            ws.merge_cells(f"B{ri}:I{ri}")
            _cell(ws, ri, 2, val)

        hr = len(info) + 3
        for col, h in enumerate([
            "Lp.", "Kompetencja", "Kategoria", "Definicja",
            "Poziom 1", "Poziom 2", "Poziom 3", "Poziom 4", "Poziom 5",
        ], 1):
            _hdr(ws, hr, col, h)

        for ci, comp in enumerate(profile.competencies, 1):
            r = hr + ci
            bg = _LIGHT_HEX if ci % 2 == 0 else None
            _cell(ws, r, 1, ci, bg=bg)
            _cell(ws, r, 2, comp.name, bold=True, bg=bg)
            _cell(ws, r, 3, CATEGORY_LABELS.get(comp.category, comp.category), bg=bg)
            _cell(ws, r, 4, comp.definition, bg=bg)
            for j, lk in enumerate(LEVEL_KEYS, 5):
                _cell(ws, r, j, comp.get_level_description(lk), bg=bg)
            ws.row_dimensions[r].height = 80

        for col, w in enumerate([4, 22, 14, 36, 28, 28, 28, 28, 28], 1):
            ws.column_dimensions[get_column_letter(col)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# OSP — Opis Stanowiska Pracy (Formularz OSP RES 1.4)
# ---------------------------------------------------------------------------
def build_osp_docx(d: dict) -> bytes:
    """Generuje DOCX Opisu Stanowiska Pracy — 13 sekcji, branding Enterprise Advisors."""
    doc = Document()
    _style_base(doc)

    def _gold_bar():
        tbl = doc.add_table(rows=1, cols=1)
        c = tbl.rows[0].cells[0]
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F5C518")
        tcPr.append(shd)
        c.paragraphs[0].paragraph_format.space_before = Pt(2)
        c.paragraphs[0].paragraph_format.space_after = Pt(2)

    def _sec(num, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        rn = p.add_run(f"{num}. "); rn.font.name = FONT; rn.font.bold = True; rn.font.size = Pt(11); rn.font.color.rgb = EA_GOLD
        rt = p.add_run(text.upper()); rt.font.name = FONT; rt.font.bold = True; rt.font.size = Pt(11); rt.font.color.rgb = EA_NAVY

    def _info_tbl(rows_data):
        if not any(v for _, v in rows_data):
            return
        tbl = doc.add_table(rows=len(rows_data), cols=2)
        tbl.style = "Table Grid"
        for i, (label, value) in enumerate(rows_data):
            c0, c1 = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
            rl = c0.paragraphs[0].add_run(label)
            rl.font.name = FONT; rl.font.size = Pt(9); rl.font.bold = True; rl.font.color.rgb = EA_NAVY
            tc = c0._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "E8EEF5")
            tcPr.append(shd)
            c0.width = Cm(5)
            rv = c1.paragraphs[0].add_run(value or "")
            rv.font.name = FONT; rv.font.size = Pt(9)
        doc.add_paragraph()

    def _hdr_tbl(headers, rows_data):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        for j, h in enumerate(headers):
            c = tbl.rows[0].cells[j]
            r = c.paragraphs[0].add_run(h)
            r.font.name = FONT; r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1B4F8A")
            tcPr.append(shd)
        for row in rows_data:
            cells = tbl.add_row().cells
            for j, v in enumerate(row):
                rv = cells[j].paragraphs[0].add_run(v or "")
                rv.font.name = FONT; rv.font.size = Pt(9)
        doc.add_paragraph()

    # --- Strona tytułowa ---
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("OPIS STANOWISKA PRACY")
    tr.font.name = FONT; tr.font.bold = True; tr.font.size = Pt(22); tr.font.color.rgb = EA_NAVY
    _gold_bar()
    doc.add_paragraph()
    if d.get("nazwa_stanowiska"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(d["nazwa_stanowiska"]); r.font.name = FONT; r.font.bold = True; r.font.size = Pt(16); r.font.color.rgb = EA_NAVY
    if d.get("jednostka"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(d["jednostka"]); r.font.name = FONT; r.font.size = Pt(12); r.font.color.rgb = EA_GRAY
    gp = doc.add_paragraph(); gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gr = gp.add_run(f"Enterprise Advisors  |  {datetime.now().strftime('%Y-%m-%d')}")
    gr.font.name = FONT; gr.font.size = Pt(9); gr.font.color.rgb = EA_GRAY
    doc.add_page_break()

    _sec("1", "Identyfikacja stanowiska")
    _info_tbl([("Numer", d.get("numer", "")), ("Nazwa stanowiska", d.get("nazwa_stanowiska", ""))])

    _sec("2", "Cel stanowiska")
    _para(doc, d.get("cel", "") or "(nie określono)")
    doc.add_paragraph()

    _sec("3", "Umiejscowienie w strukturze")
    _info_tbl([
        ("Jednostka organizacyjna", d.get("jednostka", "")),
        ("Komórka wykonawcza", d.get("komorka", "")),
        ("Stanowisko przełożonego", d.get("przelozony", "")),
        ("Stanowiska podległe", d.get("podlegle", "")),
    ])

    _sec("4", "Doświadczenie i kwalifikacje")
    _hdr_tbl(
        ["", "Wymagane", "Pożądane"],
        [
            ["Staż pracy", d.get("staz_wym", ""), d.get("staz_poz", "")],
            ["Stanowisko", d.get("stanowisko_wym", ""), d.get("stanowisko_poz", "")],
            ["Specjalność", d.get("specjalnosc_wym", ""), d.get("specjalnosc_poz", "")],
            ["Wykształcenie", d.get("wyksztalcenie_wym", ""), d.get("wyksztalcenie_poz", "")],
            ["Kwalifikacje", d.get("kwalifikacje_wym", ""), d.get("kwalifikacje_poz", "")],
            ["Uprawnienia", d.get("uprawnienia_wym", ""), d.get("uprawnienia_poz", "")],
        ],
    )

    _sec("5", "Upoważnienia i odpowiedzialność")
    _para(doc, d.get("upowaznienia", "") or "(nie określono)")
    doc.add_paragraph()

    _sec("6", "Kluczowe wskaźniki efektywności")
    _ws = [l.strip() for l in (d.get("kpi_wskazniki", "") or "").splitlines() if l.strip()]
    _kr = [l.strip() for l in (d.get("kpi_kryteria", "") or "").splitlines() if l.strip()]
    _mk = max(len(_ws), len(_kr), 4)
    _hdr_tbl(
        ["Nr", "Wskaźnik (KPI)", "Kryterium oceny skuteczności"],
        [[str(i + 1), _ws[i] if i < len(_ws) else "", _kr[i] if i < len(_kr) else ""] for i in range(_mk)],
    )

    _sec("7", "Zadania wykonywane na stanowisku")
    _zad = [l.strip() for l in (d.get("zadania", "") or "").splitlines() if l.strip()]
    if _zad:
        _hdr_tbl(["Nr", "Zadanie realizowane na stanowisku"], [[str(i + 1), z] for i, z in enumerate(_zad)])
    else:
        _para(doc, "(nie określono)"); doc.add_paragraph()

    _sec("8", "Kompetencje")
    _komps = d.get("kompetencje", [])
    if _komps:
        _hdr_tbl(
            ["Nr KK", "Kompetencja", "Poziom"],
            [[k.get("nr", str(i + 1)), k.get("nazwa", ""), k.get("poziom", "")] for i, k in enumerate(_komps)],
        )
    else:
        _para(doc, "(nie określono)"); doc.add_paragraph()

    _sec("9", "Pozostałe elementy opisu")
    _info_tbl([
        ("Wyposażenie i środki pracy", d.get("wyposazenie", "")),
        ("Uciążliwość", d.get("uciazkliwosc", "")),
        ("Zagrożenia", d.get("zagrozenia", "")),
        ("Kto zastępuje w trakcie urlopu?", d.get("zastepstwo_urlop", "")),
        ("Inne wymogi", d.get("inne_wymogi", "")),
    ])

    _sec("10", "Relacje i współpraca")
    _info_tbl([
        ("Współpraca wewnętrzna (działy, funkcje)", d.get("wspolpraca_wewn", "")),
        ("Współpraca zewnętrzna (klienci, dostawcy, partnerzy)", d.get("wspolpraca_zewn", "")),
    ])

    _sec("11", "Zastępstwa")
    _info_tbl([
        ("Kogo zastępuje to stanowisko", d.get("zastepuje_kogo", "")),
        ("Kto zastępuje to stanowisko", d.get("kto_zastepuje", "")),
    ])

    _sec("12", "Limity decyzyjne")
    _info_tbl([
        ("Decyzje podejmowane samodzielnie", d.get("decyzje_sam", "")),
        ("Decyzje wymagające akceptacji przełożonego", d.get("decyzje_akc", "")),
        ("Limity finansowe (kwoty, budżet)", d.get("limity_fin", "")),
        ("Pełnomocnictwa i upoważnienia", d.get("pelnomocnictwa", "")),
    ])

    _sec("13", "Miejsce i tryb pracy")
    _info_tbl([
        ("Miejsce wykonywania pracy", d.get("miejsce", "")),
        ("Tryb pracy (stacjonarny, hybrydowy, zdalny)", d.get("tryb", "")),
        ("Mobilność i delegacje", d.get("mobilnosc", "")),
        ("Dyspozycyjność (zmiany, dyżury)", d.get("dyspozycyjnosc", "")),
    ])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
